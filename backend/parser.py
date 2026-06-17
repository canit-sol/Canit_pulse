"""
Extracts raw text from uploaded files (PDF, DOCX, TXT).
Uses PyMuPDF for PDF (no cffi issues on Windows).
"""
import io
def extract_text_from_file(file_bytes: bytes, ext: str) -> str:
    ext = ext.lower()

    if ext == ".txt":
        return file_bytes.decode("utf-8", errors="ignore")

    elif ext in (".docx", ".doc"):
        return _extract_docx(file_bytes)

    elif ext == ".pdf":
        return _extract_pdf(file_bytes)

    return ""


def _extract_docx(file_bytes: bytes) -> str:
    try:
        import docx
        doc = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except Exception as e:
        raise ValueError(f"Could not read DOCX file: {e}")


def _extract_pdf(file_bytes: bytes) -> str:
    try:
        from pypdf import PdfReader
        import io
        
        reader = PdfReader(io.BytesIO(file_bytes))
        return "\n".join(page.extract_text() for page in reader.pages if page.extract_text())
    except Exception as e:
        raise ValueError(f"Could not read PDF file: {e}")